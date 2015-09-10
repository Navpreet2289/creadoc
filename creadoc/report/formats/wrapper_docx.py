# coding: utf-8
from copy import copy, deepcopy
from operator import attrgetter, itemgetter
from docx import Document
from creadoc.enums import SourceType
from creadoc.exceptions import CreaDocException
from creadoc.helper.tags import tag_data, get_by_key
from creadoc.registry.source import SourceRegistry
from creadoc.report.constants import (
    RE_TAG_TEMPLATE, OPEN_TAG, CLOSE_TAG,
    RE_START_CYCLE_TEMPLATE, RE_END_CYCLE_TEMPLATE,
    CLOSE_BLOCK_TAG, OPEN_BLOCK_TAG)
from creadoc.report.formats.interface import CreaDocFormatWrapper

__author__ = 'damirazo <me@damirazo.ru>'


class DocxCreaDocFormatWrapper(CreaDocFormatWrapper):
    u"""
    Класс-обертка для работы с документами в формате docx
    """

    def main(self):
        self.document = Document(self.source_path)

    def tags(self):
        u"""
        Список тегов внутри документа
        Теги, находящиеся внутри блока цикла являются динамическими
        и не возвращаются данным методом
        """
        tags = []
        block_started = False

        for paragraph in self.document.paragraphs:
            for run in paragraph.runs:
                if RE_START_CYCLE_TEMPLATE.search(run.text):
                    block_started = True

                if RE_END_CYCLE_TEMPLATE.search(run.text):
                    block_started = False

                # Учитываем только теги,
                # которые находятся за пределами блока цикла
                if not block_started and RE_TAG_TEMPLATE.search(run.text):
                    tags.append(RE_TAG_TEMPLATE.findall(run.text)[0])

        return tags

    def has_tag(self, tag_name):
        for full_tag, tag, modifier, _ in self.tags():
            if tag == tag_name:
                return True

        return False

    def sources(self):
        u"""
        Список источников данных верхнего уровня
        """
        result = {}

        for joined_tags in self.tags():
            segments = joined_tags[1].split('.')
            root_tag = segments[0]

            source = SourceRegistry.source_by_tag(root_tag)
            result[root_tag] = source

        return result

    def save(self, path):
        u"""
        Сохранение документа
        """
        self.document.save(path)

    def replace_tags(self, params):
        u"""
        Замена тегов на указанные эквиваленты в шаблоне
        """
        for paragraph in self.document.paragraphs:
            for run in paragraph.runs:
                if RE_TAG_TEMPLATE.search(run.text):
                    tag = tag_data(run.text)

                    if not self.has_tag(tag['tag_name']):
                        continue

                    run.text = run.text.replace(
                        tag['full_tag'], params.get(tag['tag_name'], ''))

    def prepare_cycles(self):
        u"""
        Обработка циклов
        """
        # Глобальное смещение индексов
        # Происходит потому, что при копировании элементов в блоке с циклом
        # нам необходимо также увеличить значения всех индексов
        global_shift_index = 0
        # Получение реального индекса на основе глобального
        index = lambda num: num + global_shift_index

        for cycle in self.cycles():
            # Параметры цикла:
            # 1. Внешнее наименование списочного тега
            # 2. Наименование тега на каждой итерации
            # для одного из объектов цикла
            full_tag, tag, inner_tag = cycle['params'][0]

            # Индекс параграфа, с которого начинается блок цикла
            begin_paragraph = index(cycle['begin_paragraph'])
            # Индекс параграфа, на котором завершается блок цикла
            end_paragraph = index(cycle['end_paragraph'])
            # Индекс рана, в котором начинается блок цикла
            begin_run = cycle['begin_run']
            # Индекс рана, в котором завершается блок цикла
            end_run = cycle['end_run']

            source = SourceRegistry.source_by_tag(tag)

            if source.type != SourceType.LIST:
                raise CreaDocException(
                    u'"{}" не является списочным тегом!'.format(source.tag))

            # Срез всех параграфов, которые входят в блок цикла
            paragraphs_in_block = self.document.paragraphs[
                begin_paragraph + 1: end_paragraph
            ]

            # Список копируемых в рамках блока параграфов
            paragraphs_to_copy = []
            paragraph_enumerator = enumerate(
                paragraphs_in_block,
                start=begin_paragraph + 1
            )

            # Сохраняем копии всех копируемых параграфов
            for i, paragraph in paragraph_enumerator:
                paragraphs_to_copy.append(deepcopy(paragraph))

            rows = source.harvest_data()
            previous_paragraph = self.document.paragraphs[end_paragraph - 1]

            for i, row in enumerate(rows, start=1):
                if i > 1:
                    for paragraph in paragraphs_to_copy:
                        copied_paragraph = deepcopy(paragraph)
                        previous_paragraph._p.addnext(copied_paragraph._p)
                        previous_paragraph = copied_paragraph

                        # Смещаем глобальный индекс
                        global_shift_index += 1

                        self._prepare_cycle_runs(
                            i,
                            row,
                            inner_tag,
                            copied_paragraph.runs
                        )

                else:
                    for paragraph in paragraphs_in_block:
                        self._prepare_cycle_runs(
                            i,
                            row,
                            inner_tag,
                            paragraph.runs
                        )

    def _prepare_cycle_runs(self, iteration_num, row, inner_tag_name, runs):
        u"""
        Обработка ранов для блока цикла

        :param iteration_num: Номер текущей итерации
        :param row: Объект записи для текущей итерации
        :param inner_tag_name: Имя списочного тега на каждой итерации
        :param runs: Список обрабатываемых ранов (в рамках параграфа)
        """
        iteration_number_tag = u'НомерИтерации'

        for run in runs:
            if RE_TAG_TEMPLATE.search(run.text):
                _tag_data = tag_data(run.text)
                _tag_data_segments = (
                    _tag_data['tag_name'].split('.'))

                root_tag = _tag_data_segments[0]

                # Тег для текущей итерации
                if root_tag == inner_tag_name:
                    value = get_by_key(
                        row,
                        _tag_data_segments[1:]
                    )
                elif root_tag == iteration_number_tag:
                    value = iteration_num
                # Обычные теги из верхней области видимости
                else:
                    # FIXME: Не забыть убрать
                    value = 'TEST'

                run.text = run.text.replace(
                    _tag_data['full_tag'],
                    unicode(value)
                )

    def normalize(self):
        u"""
        Нормализация "ранов" в документе.
        Требуется по причине того, что один тег может быть разбит сразу
        по нескольким "ранам" по различным причинам
        (элементы орфографического словаря, различные стили для символов).
        Поэтому нам необходимо найти все "раны", из которых состоит тег,
        затем объединить их всех в один "ран".
        """
        # Нормализация обычных тегов
        self._normalize_runs(RE_TAG_TEMPLATE)
        # Нормализация тегов начала блока цикла
        self._normalize_runs(RE_START_CYCLE_TEMPLATE)
        # Нормализация тегов окончания блока цикла
        self._normalize_runs(RE_END_CYCLE_TEMPLATE)

    def cycles(self):
        u"""
        Получение информации о наличии блоков с циклами в документе
        """
        # TODO: Не находит теги, перенесенные в другие параграфы
        params = []
        cycle_block_started = False
        current_tag = {}

        for p_num, paragraph in enumerate(self.document.paragraphs):
            for r_num, run in enumerate(paragraph.runs):
                text = run.text

                if RE_START_CYCLE_TEMPLATE.search(text):
                    if cycle_block_started:
                        raise CreaDocException(
                            u'Обнаружено начало блока цикла '
                            u'до завершения другого блока цикла'
                        )

                    data = RE_START_CYCLE_TEMPLATE.findall(text)

                    cycle_block_started = True
                    current_tag = {
                        'begin_paragraph': p_num,
                        'begin_run': r_num,
                        'params': data,
                    }

                if RE_END_CYCLE_TEMPLATE.search(text):
                    if not cycle_block_started:
                        raise CreaDocException(
                            u'Обнаружен закрывающий тег блока с циклом, '
                            u'однако блок с циклом еще не открыт'
                        )

                    cycle_block_started = False
                    current_tag.update({
                        'end_paragraph': p_num,
                        'end_run': r_num,
                    })
                    params.append(current_tag)

        # Проверяем, что мы закрыли все обнаруженные теги
        if current_tag is not None and cycle_block_started:
            raise CreaDocException(u'В шаблоне имеется незакрытый тег цикла')

        return params

    def _paragraphs(self):
        u"""
        Список параграфов внутри документа
        """
        result = []

        for paragraph in self.document.paragraphs:
            result.append(u''.join(map(attrgetter('text'), paragraph.runs)))

        return result

    def _full_text(self):
        u"""
        Полный текст документа
        """
        return u'\n'.join(self._paragraphs())

    def _normalize_runs(self, search_tag_template):
        u"""
        Выполнение нормализации "ранов" для тегов
        с указанными признаками открытия и закрытия
        """
        for paragraph in self.document.paragraphs:
            paragraph_texts = map(lambda x: x.text, paragraph.runs)

            params = set(map(
                itemgetter(0),
                search_tag_template.findall(paragraph.text)
            ))

            for param in params:
                current_text = u''
                begin_index = 0
                end_index = 0

                for run_index, text in enumerate(paragraph_texts):
                    current_text += text

                    if param in current_text:
                        end_index = run_index
                        break

                current_text = u''

                for run_index, text in enumerate(reversed(paragraph_texts)):
                    current_text = text + current_text

                    if param in current_text:
                        begin_index = len(paragraph_texts) - run_index - 1
                        break

                tag_value = u''

                for run_index, run in enumerate(paragraph.runs):
                    if begin_index <= run_index <= end_index:
                        tag_value += run.text

                for run_index, run in enumerate(paragraph.runs):
                    if begin_index == run_index:
                        run.text = tag_value
                    if begin_index < run_index <= end_index:
                        run.text = u''

                new_text = paragraph.text
