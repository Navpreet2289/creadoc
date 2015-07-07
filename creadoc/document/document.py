# coding: utf-8
import re
from operator import attrgetter
from docx import Document

__author__ = 'damirazo <me@damirazo.ru>'


class CreaDoc(object):
    u"""
    Объектное представление шаблона печатной формы
    """

    # Паттерн для поиска тегов в шаблоне
    tags_pattern = re.compile(
        '{{\s*([\w\d\.а-яА-ЯйЙёЁ]+?)\s*}}', re.I | re.U)

    def __init__(self, path):
        u"""
        :param path: Путь до шаблона
        """
        self.path = path
        self.document = Document(self.path)

    @property
    def paragraphs(self):
        u"""
        Список параграфов внутри документа
        """
        result = []

        for paragraph in self.document.paragraphs:
            result.append(u''.join(map(attrgetter('text'), paragraph.runs)))

        return result

    @property
    def full_text(self):
        u"""
        Полный текст документа
        """
        return u'\n'.join(self.paragraphs)

    @property
    def tags(self):
        u"""
        Список тегов внутри документа
        """
        return self.tags_pattern.findall(self.full_text)

    def save(self, path=None):
        return self.document.save(path or self.path)
